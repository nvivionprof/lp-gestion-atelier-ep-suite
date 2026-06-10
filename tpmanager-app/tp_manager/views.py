from __future__ import annotations
import csv
from io import BytesIO, TextIOWrapper
from django.conf import settings
from django.contrib import messages
from django.core.exceptions import PermissionDenied, ValidationError
from django.db.models import Count, Q
from django.http import FileResponse, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.utils import timezone
from django.views.decorators.http import require_http_methods
from django.views.decorators.csrf import csrf_exempt
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from .context_processors import current_tp_user
from .forms import (
    TPForm, TPFilterForm, TPDocumentForm, TPFormationNiveauForm, TPSystemeForm, TPCompetenceForm,
    TPPrerequisForm, TPSuivantForm, ZoneApprentissageForm, ThemeGeneralForm, ThemeSecondaireForm,
    TypeTPForm, SystemeRefForm, NiveauForm, FormationNiveauForm, SequenceForm, SequenceTPForm,
    ParcoursEleveTPForm, StudentParcoursUpdateForm, ProfParcoursUpdateForm, TraceEleveTPForm,
    EvaluationCompetenceForm, ReferentielImportCsvForm, TPTacheForm, TPSavoirForm, TPCritereForm, TPContributionPermissionForm,
)
from .models import (
    TpUser, Formation, Niveau, FormationNiveau, ZoneApprentissage, ThemeGeneral, ThemeSecondaire, TypeTP,
    SystemePedagogiqueRef, Referentiel, BlocCompetence, Competence, SousCompetence,
    ActiviteReferentiel, TacheReferentiel, TP, TPFormationNiveau, TPSysteme, TPCompetence,
    TPPrerequis, TPSuivant, TPDocument, SerieTP, SequencePedagogique, SequenceTP,
    ParcoursEleveTP, TraceEleveTP, EvaluationCompetenceTP, SavoirAssocie, CritereEvaluation, TPContributionPermission,
    TPTache, TPSavoir, TPCritere,
)
from .permissions import tp_login_required, tp_prof_required, tp_admin_required
from .sync import sync_users_from_lp_core, sync_formations_from_lp_core, sync_systems_from_system_manager


def _active_contribution_permissions(user, tp=None):
    if not user:
        return TPContributionPermission.objects.none()
    now = timezone.now()
    qs = TPContributionPermission.objects.filter(eleve=user, actif=True, date_debut__lte=now).filter(Q(date_fin__isnull=True) | Q(date_fin__gte=now))
    if tp is not None:
        qs = qs.filter(Q(tp__isnull=True) | Q(tp=tp))
    return qs


def _can_student_create_tp(user):
    return bool(user and not user.is_prof_like and _active_contribution_permissions(user).filter(peut_creer=True).exists())


def _can_edit_tp(user, tp):
    if not user:
        return False
    if user.is_prof_like:
        return True
    if tp.auteur_id != user.id or tp.statut not in {'brouillon', 'brouillon_eleve', 'relecture'}:
        return False
    return _active_contribution_permissions(user, tp=tp).filter(peut_modifier=True).exists()


def _can_add_doc_to_tp(user, tp):
    if not user:
        return False
    if user.is_prof_like:
        return True
    if tp.auteur_id != user.id:
        return False
    return _active_contribution_permissions(user, tp=tp).filter(peut_ajouter_documents=True).exists()





# --- LP Suite SSO portal-login V0.1.0 ---
def _portal_token_payload(request):
    from django.core import signing
    try:
        return signing.loads(
            request.GET.get('token') or '',
            key=getattr(settings, 'LP_CORE_API_TOKEN', ''),
            salt='lp-suite-sso',
            max_age=120,
        )
    except Exception:
        return None

def portal_login(request):
    payload = _portal_token_payload(request)
    if not payload:
        messages.error(request, 'Connexion LP Core impossible ou expirée. Merci de te reconnecter.')
        return redirect('tp_login')
    code = (payload.get('code') or '').strip()
    username = (payload.get('username') or '').strip()
    user = TpUser.objects.filter(Q(code=code) | Q(username=username), active=True).first()
    if not user:
        messages.error(request, 'Compte LP Core non synchronisé dans TP Manager.')
        return redirect('tp_login')
    request.session['tp_user_id'] = user.id
    messages.success(request, f'Connexion TP Manager via LP Core : {user.full_name}.')
    return redirect('tp_dashboard')

def login_view(request):
    if request.method == 'POST':
        username = (request.POST.get('username') or '').strip()
        password = request.POST.get('password') or ''
        user = TpUser.objects.filter(username=username, active=True).first()
        if user and user.check_password(password):
            request.session['tp_user_id'] = user.id
            messages.success(request, f'Connexion TP Manager : {user.full_name}.')
            return redirect('tp_dashboard')
        messages.error(request, 'Identifiant ou mot de passe incorrect.')
    return render(request, 'tp_manager/login.html')


def logout_view(request):
    request.session.pop('tp_user_id', None)
    messages.success(request, 'Déconnexion TP Manager effectuée.')
    return redirect('tp_dashboard')


def dashboard(request):
    user = current_tp_user(request)
    context = {
        'tp_count': TP.objects.count(),
        'published_count': TP.objects.filter(statut='publie').count(),
        'sequence_count': SequencePedagogique.objects.exclude(statut='archive').count(),
        'parcours_count': ParcoursEleveTP.objects.exclude(statut='archive').count(),
        'active_contrib_count': TPContributionPermission.objects.filter(actif=True).count(),
        'recent_tps': TP.objects.select_related('formation_principale', 'zone_apprentissage', 'theme_secondaire').order_by('-updated_at')[:8],
        'recent_parcours': ParcoursEleveTP.objects.select_related('eleve', 'tp').order_by('-updated_at')[:8] if user and user.is_prof_like else [],
    }
    if user and not user.is_prof_like:
        context['my_parcours'] = ParcoursEleveTP.objects.select_related('tp').filter(eleve=user).order_by('-updated_at')[:8]
        context['can_contribute_tp'] = _can_student_create_tp(user)
    return render(request, 'tp_manager/dashboard.html', context)


def _internal_sync_authorized(request):
    expected = getattr(settings, 'LP_CORE_API_TOKEN', '') or ''
    provided = request.headers.get('X-API-Key') or request.POST.get('token') or request.GET.get('token') or ''
    return bool(expected and provided == expected)


@csrf_exempt
@require_http_methods(['POST'])
def sync_lp_core_view(request):
    internal = _internal_sync_authorized(request)
    if internal:
        try:
            force_password = request.POST.get('force_password') in {'1', 'true', 'True', 'oui', 'OUI'}
            core_user_id = request.POST.get('core_user_id') or request.GET.get('core_user_id')
            users_report = sync_users_from_lp_core(timeout=90, force_password=force_password, core_user_id=core_user_id)
            forms_report = {'created': 0, 'updated': 0, 'errors': []}
            if not core_user_id:
                forms_report = sync_formations_from_lp_core(timeout=90)
            errors = users_report.get('errors', []) + forms_report.get('errors', [])
            return JsonResponse({'ok': len(errors) == 0, 'users': users_report, 'formations': forms_report, 'errors': errors})
        except Exception as exc:
            return JsonResponse({'ok': False, 'error': str(exc)}, status=500)
    user = current_tp_user(request)
    if not user or not user.is_admin_like:
        messages.error(request, 'Synchronisation réservée aux administrateurs.')
        return redirect('tp_login')
    try:
        users_report = sync_users_from_lp_core(timeout=90)
        forms_report = sync_formations_from_lp_core(timeout=90)
        messages.success(request, f"Synchro LP Core terminée : utilisateurs {users_report.get('created',0)} créés / {users_report.get('updated',0)} mis à jour, formations {forms_report.get('created',0)} créées / {forms_report.get('updated',0)} mises à jour.")
    except Exception as exc:
        messages.error(request, f'Échec synchronisation LP Core : {exc}')
    return redirect('tp_referentials')


@require_http_methods(['POST'])
def sync_systems_view(request):
    user = current_tp_user(request)
    if not user or not user.is_admin_like:
        messages.error(request, 'Synchronisation systèmes réservée aux administrateurs.')
        return redirect('tp_login')
    try:
        report = sync_systems_from_system_manager()
        messages.success(request, f"Synchro System Manager terminée : {report.get('created',0)} créés / {report.get('updated',0)} mis à jour.")
    except Exception as exc:
        messages.error(request, f'Échec synchronisation systèmes : {exc}')
    return redirect('tp_referentials')


def _filter_tps(request, student_mode=False):
    form = TPFilterForm(request.GET or None)
    tps = TP.objects.select_related('formation_principale', 'zone_apprentissage', 'theme_general', 'theme_secondaire', 'type_tp')
    if student_mode:
        tps = tps.filter(statut='publie')
    if form.is_valid():
        cd = form.cleaned_data
        q = cd.get('q')
        if q:
            tps = tps.filter(Q(code__icontains=q) | Q(titre__icontains=q) | Q(resume_apprentissages__icontains=q))
        if cd.get('formation'):
            tps = tps.filter(Q(formation_principale=cd['formation']) | Q(formations_niveaux__formation=cd['formation']))
        if cd.get('niveau'):
            tps = tps.filter(formations_niveaux__niveau=cd['niveau'])
        if cd.get('zone'):
            tps = tps.filter(zone_apprentissage=cd['zone'])
        if cd.get('theme_general'):
            tps = tps.filter(theme_general=cd['theme_general'])
        if cd.get('theme_secondaire'):
            tps = tps.filter(theme_secondaire=cd['theme_secondaire'])
        if cd.get('competence'):
            tps = tps.filter(competences__competence=cd['competence'])
        if cd.get('temps_max'):
            tps = tps.filter(temps_estime_minutes__lte=cd['temps_max'])
    return form, tps.distinct().order_by('code')


def tp_list(request):
    user = current_tp_user(request)
    student_mode = bool(user and not user.is_prof_like)
    form, tps = _filter_tps(request, student_mode=student_mode)
    return render(request, 'tp_manager/tp_list.html', {'form': form, 'tps': tps[:500], 'student_mode': student_mode})


@tp_login_required
def tp_create(request):
    user = current_tp_user(request)
    if not (user and (user.is_prof_like or _can_student_create_tp(user))):
        messages.error(request, 'Création de TP réservée aux professeurs ou aux élèves disposant d’un droit temporaire actif.')
        return redirect('tp_dashboard')
    form = TPForm(request.POST or None, request.FILES or None, allow_status=user.is_prof_like)
    if request.method == 'POST' and form.is_valid():
        tp = form.save(commit=False)
        tp.auteur = user
        if not user.is_prof_like:
            tp.statut = 'brouillon_eleve'
        tp.save()
        messages.success(request, f'TP créé : {tp.code}.')
        return redirect('tp_detail', tp.pk)
    title = 'Créer un TP modèle' if user.is_prof_like else 'Créer un brouillon de TP élève'
    return render(request, 'tp_manager/form.html', {'form': form, 'title': title})


def _visible_documents(tp, user):
    docs = tp.documents.filter(actif=True)
    if user and user.is_prof_like:
        return docs.filter(visible_prof=True)
    return docs.filter(visible_eleve=True)


def tp_detail(request, pk):
    tp = get_object_or_404(TP.objects.select_related('formation_principale', 'zone_apprentissage', 'theme_general', 'theme_secondaire', 'type_tp', 'auteur'), pk=pk)
    user = current_tp_user(request)
    if tp.statut != 'publie' and not (user and (user.is_prof_like or _can_edit_tp(user, tp))):
        messages.error(request, 'TP non publié.')
        return redirect('tp_list')
    parcours = None
    missing = []
    if user and not user.is_prof_like:
        parcours = ParcoursEleveTP.objects.filter(eleve=user, tp=tp).first()
        if parcours:
            missing = parcours.prerequis_valides()
    return render(request, 'tp_manager/tp_detail.html', {
        'tp': tp,
        'documents': _visible_documents(tp, user),
        'parcours': parcours,
        'missing_prerequis': missing,
    })


@tp_login_required
def tp_update(request, pk):
    tp = get_object_or_404(TP, pk=pk)
    user = current_tp_user(request)
    if not _can_edit_tp(user, tp):
        messages.error(request, 'Modification non autorisée pour ce TP.')
        return redirect('tp_detail', tp.pk)
    form = TPForm(request.POST or None, request.FILES or None, instance=tp, allow_status=user.is_prof_like)
    if request.method == 'POST' and form.is_valid():
        obj = form.save(commit=False)
        if not user.is_prof_like:
            obj.statut = 'brouillon_eleve'
        obj.save()
        messages.success(request, 'TP modifié.')
        return redirect('tp_detail', tp.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Modifier {tp.code}'})


@tp_login_required
def document_add(request, tp_pk):
    tp = get_object_or_404(TP, pk=tp_pk)
    user = current_tp_user(request)
    if not _can_add_doc_to_tp(user, tp):
        messages.error(request, 'Ajout de document non autorisé pour ce TP.')
        return redirect('tp_detail', tp.pk)
    form = TPDocumentForm(request.POST or None, request.FILES or None)
    if request.method == 'POST' and form.is_valid():
        doc = form.save(commit=False)
        doc.tp = tp
        doc.uploaded_by = user
        if not user.is_prof_like:
            doc.visible_eleve = False
            doc.visible_prof = True
        doc.save()
        messages.success(request, 'Document ajouté.')
        return redirect('tp_detail', tp.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Ajouter un document — {tp.code}'})


@tp_prof_required
def tpformation_add(request, tp_pk):
    tp = get_object_or_404(TP, pk=tp_pk)
    form = TPFormationNiveauForm(request.POST or None)
    if request.method == 'POST' and form.is_valid():
        item = form.save(commit=False)
        item.tp = tp
        item.save()
        messages.success(request, 'Formation/niveau ajouté.')
        return redirect('tp_detail', tp.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Ajouter une formation/niveau — {tp.code}'})


@tp_prof_required
def tpsysteme_add(request, tp_pk):
    tp = get_object_or_404(TP, pk=tp_pk)
    form = TPSystemeForm(request.POST or None)
    if request.method == 'POST' and form.is_valid():
        item = form.save(commit=False)
        item.tp = tp
        item.save()
        messages.success(request, 'Système associé.')
        return redirect('tp_detail', tp.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Associer un système — {tp.code}'})


@tp_prof_required
def tpcompetence_add(request, tp_pk):
    tp = get_object_or_404(TP, pk=tp_pk)
    form = TPCompetenceForm(request.POST or None, formation=tp.formation_principale)
    if request.method == 'POST' and form.is_valid():
        item = form.save(commit=False)
        item.tp = tp
        item.save()
        messages.success(request, 'Compétence associée.')
        return redirect('tp_detail', tp.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Associer une compétence — {tp.code}'})



@tp_prof_required
def tptache_add(request, tp_pk):
    tp = get_object_or_404(TP, pk=tp_pk)
    form = TPTacheForm(request.POST or None, formation=tp.formation_principale)
    if request.method == 'POST' and form.is_valid():
        item = form.save(commit=False)
        item.tp = tp
        item.save()
        messages.success(request, 'Tâche officielle associée.')
        return redirect('tp_detail', tp.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Associer une tâche officielle — {tp.code}'})


@tp_prof_required
def tpsavoir_add(request, tp_pk):
    tp = get_object_or_404(TP, pk=tp_pk)
    form = TPSavoirForm(request.POST or None, formation=tp.formation_principale)
    if request.method == 'POST' and form.is_valid():
        item = form.save(commit=False)
        item.tp = tp
        item.save()
        messages.success(request, 'Savoir associé au TP.')
        return redirect('tp_detail', tp.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Associer un savoir — {tp.code}'})


@tp_prof_required
def tpcritere_add(request, tp_pk):
    tp = get_object_or_404(TP, pk=tp_pk)
    form = TPCritereForm(request.POST or None, formation=tp.formation_principale)
    if request.method == 'POST' and form.is_valid():
        item = form.save(commit=False)
        item.tp = tp
        item.save()
        messages.success(request, 'Critère d’évaluation associé au TP.')
        return redirect('tp_detail', tp.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Associer un critère — {tp.code}'})


@tp_prof_required
def tpprerequis_add(request, tp_pk):
    tp = get_object_or_404(TP, pk=tp_pk)
    form = TPPrerequisForm(request.POST or None, current_tp=tp)
    if request.method == 'POST' and form.is_valid():
        item = form.save(commit=False)
        item.tp = tp
        item.full_clean()
        item.save()
        messages.success(request, 'Prérequis ajouté.')
        return redirect('tp_detail', tp.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Ajouter un prérequis — {tp.code}'})


@tp_prof_required
def tpsuivant_add(request, tp_pk):
    tp = get_object_or_404(TP, pk=tp_pk)
    form = TPSuivantForm(request.POST or None, current_tp=tp)
    if request.method == 'POST' and form.is_valid():
        item = form.save(commit=False)
        item.tp = tp
        item.save()
        messages.success(request, 'TP suivant conseillé ajouté.')
        return redirect('tp_detail', tp.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Ajouter un TP suivant — {tp.code}'})


@tp_login_required
def request_tp_for_me(request, pk):
    user = current_tp_user(request)
    tp = get_object_or_404(TP, pk=pk, statut='publie')
    parcours, created = ParcoursEleveTP.objects.get_or_create(eleve=user, tp=tp, sequence=None, defaults={'statut': 'a_faire'})
    missing = parcours.prerequis_valides()
    if missing:
        parcours.statut = 'bloque'
        parcours.save(update_fields=['statut', 'updated_at'])
        messages.warning(request, 'TP ajouté mais bloqué : prérequis non validés.')
    else:
        if parcours.statut == 'bloque':
            parcours.statut = 'a_faire'
            parcours.save(update_fields=['statut', 'updated_at'])
        messages.success(request, 'TP ajouté à ton parcours.' if created else 'TP déjà présent dans ton parcours.')
    return redirect('parcours_detail', parcours.pk)


@tp_login_required
def parcours_list(request):
    user = current_tp_user(request)
    qs = ParcoursEleveTP.objects.select_related('eleve', 'tp', 'sequence', 'systeme_utilise')
    if not user.is_prof_like:
        qs = qs.filter(eleve=user)
    q = (request.GET.get('q') or '').strip()
    if q:
        qs = qs.filter(Q(tp__code__icontains=q) | Q(tp__titre__icontains=q) | Q(eleve__last_name__icontains=q) | Q(eleve__first_name__icontains=q))
    return render(request, 'tp_manager/parcours_list.html', {'parcours_list': qs.order_by('-updated_at')[:700], 'q': q})


@tp_login_required
def parcours_detail(request, pk):
    parcours = get_object_or_404(ParcoursEleveTP.objects.select_related('eleve', 'tp', 'sequence', 'systeme_utilise'), pk=pk)
    user = current_tp_user(request)
    if not user.is_prof_like and parcours.eleve_id != user.id:
        raise PermissionDenied
    student_form = StudentParcoursUpdateForm(request.POST or None, instance=parcours, prefix='student')
    prof_form = ProfParcoursUpdateForm(request.POST or None, instance=parcours, prefix='prof') if user.is_prof_like else None
    if request.method == 'POST':
        if 'student-submit' in request.POST and student_form.is_valid():
            obj = student_form.save(commit=False)
            if obj.statut == 'en_cours' and not obj.date_debut:
                obj.date_debut = timezone.now()
            if obj.statut in {'realise', 'a_corriger'} and not obj.date_realisation:
                obj.date_realisation = timezone.now()
            obj.save()
            messages.success(request, 'Parcours mis à jour.')
            return redirect('parcours_detail', parcours.pk)
        if user.is_prof_like and 'prof-submit' in request.POST and prof_form and prof_form.is_valid():
            obj = prof_form.save(commit=False)
            if obj.statut == 'valide':
                obj.validateur = user
                obj.date_validation = timezone.now()
            obj.save()
            messages.success(request, 'Suivi professeur mis à jour.')
            return redirect('parcours_detail', parcours.pk)
    return render(request, 'tp_manager/parcours_detail.html', {
        'parcours': parcours,
        'student_form': student_form,
        'prof_form': prof_form,
        'missing_prerequis': parcours.prerequis_valides(),
    })


@tp_login_required
def trace_add(request, parcours_pk):
    parcours = get_object_or_404(ParcoursEleveTP, pk=parcours_pk)
    user = current_tp_user(request)
    if not user.is_prof_like and parcours.eleve_id != user.id:
        raise PermissionDenied
    form = TraceEleveTPForm(request.POST or None, request.FILES or None)
    if request.method == 'POST' and form.is_valid():
        trace = form.save(commit=False)
        trace.parcours = parcours
        trace.save()
        messages.success(request, 'Trace ajoutée.')
        return redirect('parcours_detail', parcours.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Ajouter une trace — {parcours.tp.code}'})


@tp_prof_required
def evaluation_add(request, parcours_pk):
    parcours = get_object_or_404(ParcoursEleveTP, pk=parcours_pk)
    form = EvaluationCompetenceForm(request.POST or None, parcours=parcours)
    if request.method == 'POST' and form.is_valid():
        ev = form.save(commit=False)
        ev.parcours = parcours
        ev.evaluateur = current_tp_user(request)
        ev.save()
        messages.success(request, 'Évaluation enregistrée.')
        return redirect('parcours_detail', parcours.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Évaluer une compétence — {parcours.eleve.full_name}'})


@tp_login_required
def parcours_export_docx(request, pk):
    parcours = get_object_or_404(ParcoursEleveTP.objects.select_related('eleve', 'tp', 'sequence', 'systeme_utilise'), pk=pk)
    user = current_tp_user(request)
    if not user.is_prof_like and parcours.eleve_id != user.id:
        raise PermissionDenied
    doc = Document()
    doc.add_heading(f'Compte rendu brut — {parcours.tp.code}', 0)
    doc.add_paragraph(f'Élève : {parcours.eleve.full_name}')
    doc.add_paragraph(f'Classe / groupe : {parcours.eleve.class_name} {parcours.eleve.group_name}')
    doc.add_paragraph(f'TP : {parcours.tp.titre}')
    doc.add_paragraph(f'Statut : {parcours.get_statut_display()}')
    if parcours.systeme_utilise:
        doc.add_paragraph(f'Système utilisé : {parcours.systeme_utilise.code} — {parcours.systeme_utilise.designation}')
    doc.add_heading('Résumé des apprentissages', level=1)
    doc.add_paragraph(parcours.tp.resume_apprentissages or 'Non renseigné.')
    doc.add_heading('Compétences travaillées', level=1)
    for link in parcours.tp.competences.select_related('competence'):
        doc.add_paragraph(f'{link.competence.code} — {link.competence.libelle} ({link.get_type_lien_display()})', style='List Bullet')
    doc.add_heading('Commentaire élève', level=1)
    doc.add_paragraph(parcours.commentaire_eleve or '')
    doc.add_heading('Traces et réponses', level=1)
    for trace in parcours.traces.all():
        doc.add_heading(trace.titre or trace.get_type_trace_display(), level=2)
        if trace.contenu_texte:
            doc.add_paragraph(trace.contenu_texte)
        if trace.fichier:
            doc.add_paragraph(f'Fichier joint : {trace.fichier.name}')
            try:
                if trace.fichier.name.lower().endswith(('.png', '.jpg', '.jpeg')):
                    doc.add_picture(trace.fichier.path, width=Inches(5.5))
            except Exception:
                pass
    doc.add_heading('Commentaire professeur', level=1)
    doc.add_paragraph(parcours.commentaire_prof or '')
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f'compte-rendu-{parcours.tp.code}-{parcours.eleve.code}.docx'.replace(' ', '_')
    return FileResponse(buffer, as_attachment=True, filename=filename)


@tp_prof_required
def sequence_list(request):
    qs = SequencePedagogique.objects.select_related('professeur', 'formation', 'niveau').order_by('-date_debut', 'titre')
    return render(request, 'tp_manager/sequence_list.html', {'sequences': qs[:500]})


@tp_prof_required
def sequence_create(request):
    form = SequenceForm(request.POST or None)
    if request.method == 'POST' and form.is_valid():
        seq = form.save(commit=False)
        seq.professeur = current_tp_user(request)
        seq.save()
        form.save_m2m()
        messages.success(request, 'Séquence créée.')
        return redirect('sequence_detail', seq.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': 'Créer une séquence'})


@tp_prof_required
def sequence_detail(request, pk):
    seq = get_object_or_404(SequencePedagogique.objects.select_related('professeur', 'formation', 'niveau'), pk=pk)
    return render(request, 'tp_manager/sequence_detail.html', {'sequence': seq})


@tp_prof_required
def sequence_add_tp(request, sequence_pk):
    seq = get_object_or_404(SequencePedagogique, pk=sequence_pk)
    form = SequenceTPForm(request.POST or None)
    if request.method == 'POST' and form.is_valid():
        item = form.save(commit=False)
        item.sequence = seq
        item.save()
        messages.success(request, 'TP ajouté à la séquence.')
        return redirect('sequence_detail', seq.pk)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': f'Ajouter un TP — {seq.titre}'})


@tp_prof_required
def sequence_assign(request, pk):
    seq = get_object_or_404(SequencePedagogique, pk=pk)
    created = 0
    for eleve in seq.eleves.filter(active=True):
        for item in seq.items.select_related('tp'):
            _, is_created = ParcoursEleveTP.objects.get_or_create(eleve=eleve, tp=item.tp, sequence=seq, defaults={'statut': 'a_faire'})
            if is_created:
                created += 1
    messages.success(request, f'Affectation terminée : {created} parcours créés.')
    return redirect('sequence_detail', seq.pk)


@tp_prof_required
def export_selection_pdf(request):
    form, tps = _filter_tps(request, student_mode=False)
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, title='Sélection TP Manager')
    styles = getSampleStyleSheet()
    story = [Paragraph('Sélection de TP disponibles', styles['Title']), Spacer(1, 12)]
    data = [['Code', 'Titre', 'Formation', 'Temps', 'Compétences']]
    for tp in tps[:120]:
        comps = ', '.join(tp.competences.select_related('competence').values_list('competence__code', flat=True)[:8])
        data.append([tp.code, tp.titre, tp.formation_principale.code if tp.formation_principale else '', f'{tp.temps_estime_minutes} min', comps])
    table = Table(data, repeatRows=1, colWidths=[75, 170, 65, 45, 145])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0b2d4d')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    story.append(table)
    doc.build(story)
    buffer.seek(0)
    return FileResponse(buffer, as_attachment=True, filename='selection-tp-manager.pdf')


@tp_admin_required
def referentials(request):
    return render(request, 'tp_manager/referentials.html', {
        'formations': Formation.objects.all(),
        'niveaux': Niveau.objects.all(),
        'zones': ZoneApprentissage.objects.all(),
        'themes_generaux': ThemeGeneral.objects.all(),
        'themes_secondaires': ThemeSecondaire.objects.select_related('theme_general'),
        'types_tp': TypeTP.objects.all(),
        'systemes': SystemePedagogiqueRef.objects.all()[:300],
        'referentiels': Referentiel.objects.select_related('formation').annotate(nb_comp=Count('blocs__competences')).all(),
        'savoirs_associes': SavoirAssocie.objects.select_related('referentiel', 'referentiel__formation').all()[:200],
        'criteres_evaluation': CritereEvaluation.objects.select_related('competence').all()[:200],
        'contributions': TPContributionPermission.objects.select_related('eleve', 'tp', 'accordee_par').all()[:200],
    })


def _simple_create_view(request, form_class, title, redirect_name='tp_referentials'):
    form = form_class(request.POST or None)
    if request.method == 'POST' and form.is_valid():
        form.save()
        messages.success(request, 'Élément enregistré.')
        return redirect(redirect_name)
    return render(request, 'tp_manager/form.html', {'form': form, 'title': title})


@tp_admin_required
def zone_create(request):
    return _simple_create_view(request, ZoneApprentissageForm, 'Créer une zone d’apprentissage')


@tp_admin_required
def theme_general_create(request):
    return _simple_create_view(request, ThemeGeneralForm, 'Créer un thème général')


@tp_admin_required
def theme_secondaire_create(request):
    return _simple_create_view(request, ThemeSecondaireForm, 'Créer un thème secondaire')


@tp_admin_required
def type_tp_create(request):
    return _simple_create_view(request, TypeTPForm, 'Créer un type de TP')


@tp_admin_required
def systeme_ref_create(request):
    return _simple_create_view(request, SystemeRefForm, 'Créer une référence système')


@tp_admin_required
def niveau_create(request):
    return _simple_create_view(request, NiveauForm, 'Créer un niveau')


@tp_admin_required
def formation_niveau_create(request):
    return _simple_create_view(request, FormationNiveauForm, 'Associer un niveau à une formation')


@tp_admin_required
@require_http_methods(['GET', 'POST'])
def referentiel_import_csv(request):
    form = ReferentielImportCsvForm(request.POST or None, request.FILES or None)
    report = None
    if request.method == 'POST' and form.is_valid():
        cd = form.cleaned_data
        ref = Referentiel.objects.create(formation=cd['formation'], nom=cd['nom'], version=cd.get('version') or '', source='import CSV')
        wrapper = TextIOWrapper(cd['fichier'].file, encoding='utf-8-sig')
        reader = csv.DictReader(wrapper, delimiter=';')
        counts = {'blocs': 0, 'competences': 0, 'sous_competences': 0, 'activites': 0, 'taches': 0}
        for row in reader:
            bcode = (row.get('bloc_code') or 'BLOC').strip()
            bloc, created = BlocCompetence.objects.get_or_create(referentiel=ref, code=bcode, defaults={'libelle': row.get('bloc_libelle') or bcode})
            counts['blocs'] += int(created)
            ccode = (row.get('competence_code') or '').strip()
            comp = None
            if ccode:
                comp, created = Competence.objects.get_or_create(bloc=bloc, code=ccode, defaults={'libelle': row.get('competence_libelle') or ccode})
                counts['competences'] += int(created)
            scode = (row.get('sous_competence_code') or '').strip()
            if comp and scode:
                _, created = SousCompetence.objects.get_or_create(competence=comp, code=scode, defaults={'libelle': row.get('sous_competence_libelle') or scode})
                counts['sous_competences'] += int(created)
            acode = (row.get('activite_code') or '').strip()
            act = None
            if acode:
                act, created = ActiviteReferentiel.objects.get_or_create(referentiel=ref, code=acode, defaults={'libelle': row.get('activite_libelle') or acode})
                counts['activites'] += int(created)
            tcode = (row.get('tache_code') or '').strip()
            if act and tcode:
                _, created = TacheReferentiel.objects.get_or_create(activite=act, code=tcode, defaults={'libelle': row.get('tache_libelle') or tcode})
                counts['taches'] += int(created)
        report = counts
        messages.success(request, f'Référentiel importé : {counts}.')
        return redirect('tp_referentials')
    return render(request, 'tp_manager/form.html', {'form': form, 'title': 'Importer un référentiel CSV', 'report': report})



@tp_prof_required
def contribution_permission_create(request):
    form = TPContributionPermissionForm(request.POST or None)
    if request.method == 'POST' and form.is_valid():
        perm = form.save(commit=False)
        perm.accordee_par = current_tp_user(request)
        perm.save()
        messages.success(request, 'Droit temporaire de contribution élève enregistré.')
        return redirect('tp_referentials')
    return render(request, 'tp_manager/form.html', {'form': form, 'title': 'Donner un droit temporaire de contribution TP'})


def api_health(request):
    return JsonResponse({'status': 'ok', 'service': 'tp-manager', 'version': settings.TPMANAGER_VERSION})


# --- Administration SQL base module ---
def _tp_sql_admin_user(request):
    from .context_processors import current_tp_user
    user = current_tp_user(request)
    return user if user and user.is_admin_like else None


def sql_database_admin(request):
    from django.contrib import messages
    from django.shortcuts import redirect
    from .db_sql_admin import render_sql_admin
    if not _tp_sql_admin_user(request):
        messages.error(request, 'Accès réservé administrateur TP Manager.')
        return redirect('tp_dashboard')
    return render_sql_admin(request, 'tp_manager/sql_database.html', 'TP Manager / Evaluation Manager')


def sql_database_export(request):
    from django.contrib import messages
    from django.shortcuts import redirect
    from .db_sql_admin import export_sql_response
    if not _tp_sql_admin_user(request):
        messages.error(request, 'Accès réservé administrateur TP Manager.')
        return redirect('tp_dashboard')
    return export_sql_response(request, 'tpmanager')


def sql_database_import(request):
    from django.contrib import messages
    from django.shortcuts import redirect
    from .db_sql_admin import import_sql_response
    if not _tp_sql_admin_user(request):
        messages.error(request, 'Accès réservé administrateur TP Manager.')
        return redirect('tp_dashboard')
    return import_sql_response(request, 'tp_manager/sql_database.html', 'TP Manager / Evaluation Manager', 'tpmanager')

def help_view(request):
    return render(request, 'tp_manager/help.html')


def about_view(request):
    return render(request, 'tp_manager/about.html')


@require_http_methods(['GET', 'POST'])
def export_pdf_config(request):
    if request.method == 'POST':
        request.session['tp_pdf_identity_mode'] = request.POST.get('identity_mode') or 'anonymous'
        messages.success(request, 'Configuration export PDF TP Manager enregistrée pour la session.')
        return redirect('tp_export_config')
    return render(request, 'tp_manager/export_pdf_config.html', {'identity_mode': request.session.get('tp_pdf_identity_mode', 'anonymous')})
